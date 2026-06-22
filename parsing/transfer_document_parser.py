"""
Transfer document parser — rule-based extraction + LLM fallback.

Parses the system-generated Ders Transfer Formu format (10-column table):
  Partner Code | Partner Name | Local Grade | ECTS Grade | Partner ECTS |
  Home Code | Home Name | ECTS Grade | IKU Grade | Home ECTS
"""

import os
import re
import logging
from datetime import datetime, timezone

from sqlalchemy.orm import Session
from db.models import (
    University, GradingScheme, GradingSchemeSnapshot, EctsIkuSnapshot,
    TransferDocument, TransferVerificationResult,
)
from backend.services.versioning import get_latest_scheme_snapshot_id, get_latest_ects_iku_snapshot_id
from matching.grade_transfer import convert_single_grade_from_snapshot
from matching.prompt_loader import load_prompt
from parsing.pdf_utils import extract_text_with_page_markers
from parsing.llm_client import parsing_llm_client

logger = logging.getLogger(__name__)

MIN_TEXT_LENGTH = 40  # below this, trigger OCR
LLM_TEXT_LIMIT = 20000  # max chars to send to LLM


def extract_text_from_file(file_path: str) -> str:
    """Extract plain text from PDF or DOCX file."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_text_with_page_markers(file_path)

    if ext in (".doc", ".docx"):
        from docx import Document
        doc = Document(file_path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(" | ".join(cells))
        return "\n".join(parts)

    raise ValueError(f"Unsupported file type: {ext}")


def _ocr_file(file_path: str) -> str:
    """OCR a textless file via multimodal LLM. Returns combined page text."""
    from parsing.pdf_utils import ocr_pdf_via_images
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return ocr_pdf_via_images(file_path)

    # DOCX can't really be image-only in practice, but handle gracefully
    if ext in (".doc", ".docx"):
        text = extract_text_from_file(file_path)
        if len(text.strip()) >= MIN_TEXT_LENGTH:
            return text
        raise ValueError("DOCX has no extractable text and OCR is not supported for DOCX files")

    raise ValueError(f"OCR not supported for file type: {ext}")


# ── Rule-based parsing ────────────────────────────────────────────────────────

_student_name_re = re.compile(
    r'(?:Öğrenci\s*(?:Ad[ıi]|İsmi)|Student\s*Name|Ad[ıi]\s*Soyad[ıi]|Name\s*Surname)\s*[:.]\s*(.+?)(?:\n|$)',
    re.IGNORECASE,
)
_student_number_re = re.compile(
    r'(?:Öğrenci\s*(?:No|Numara|Numaras[ıi]|Number|ID)|Student\s*(?:No|Number|ID))\s*[:.]\s*(\d{9,10})',
    re.IGNORECASE,
)
_fallback_number_re = re.compile(r'\b(\d{9,10})\b')

_university_line_re = re.compile(
    r'^(.+?)\s{4,}(İstanbul\s*Kültür\s*Üniversitesi|İKÜ)\s*$',
    re.IGNORECASE | re.MULTILINE,
)

# 10-column table row: partnerCode | partnerName | localGrade | ectsGrade |
# partnerEcts | homeCode | homeName | ectsGrade2 | ikuGrade | homeEcts
# Two variants:
#   full  — partnerCode + homeCode present (e.g., COM course matching)
#   no_home — homeCode+homeName empty (direct transfer), partnerCode optional

_row_full_re = re.compile(
    r'([A-Za-z0-9][A-Za-z0-9\-_./]{0,12})\s{2,}'          # partnerCode (may be empty in rare cases)
    r'(.+?)\s{2,}'                                          # partnerName
    r'([0-9,.]+|[A-Fa-f][+\-]?)\s{2,}'                     # localGrade
    r'([A-Fa-f][+\-]?)\s{2,}'                               # ectsGrade
    r'(\d+\.?\d*)\s{2,}'                                     # partnerEcts
    r'([A-Za-z0-9]{1,10})\s{2,}'                           # homeCode (e.g., COM6032)
    r'(.+?)\s{2,}'                                           # homeName
    r'([A-Fa-f][+\-]?)\s{2,}'                               # ectsGrade (repeat)
    r'([A-Fa-f][+\-]?)\s{2,}'                               # ikuGrade
    r'(\d+\.?\d*)'                                           # homeEcts
)

_row_no_home_re = re.compile(
    r'([A-Za-z0-9][A-Za-z0-9\-_./]*|)\s{2,}'               # partnerCode (optional, may be empty)
    r'(.+?)\s{2,}'                                          # partnerName
    r'([0-9,.]+|[A-Fa-f][+\-]?)\s{2,}'                     # localGrade
    r'([A-Fa-f][+\-]?)\s{2,}'                               # ectsGrade
    r'(\d+\.?\d*)\s{2,}'                                     # partnerEcts
    r'([A-Fa-f][+\-]?)\s{2,}'                               # ectsGrade (repeat)
    r'([A-Fa-f][+\-]?)\s{2,}'                               # ikuGrade
    r'(\d+\.?\d*)'                                           # homeEcts
)

_summary_re = re.compile(r'(?:Toplam|Total)\s+(\d+\.?\d*)\s+(\d+\.?\d*)', re.IGNORECASE)


def _parse_rule_based(text: str) -> dict | None:
    """Attempt rule-based extraction from a system-generated transfer form.

    Returns dict with studentName, studentNumber, partnerUniversity,
    homeUniversity, rows or None if confidence too low.
    """
    lines = text.split("\n")
    result = {
        "studentName": "",
        "studentNumber": "",
        "partnerUniversity": "",
        "homeUniversity": "İstanbul Kültür Üniversitesi",
        "rows": [],
    }

    # Student name
    m = _student_name_re.search(text)
    if m:
        name = m.group(1).strip()
        if len(name) >= 3:
            result["studentName"] = name

    # Student number
    m = _student_number_re.search(text)
    if m:
        result["studentNumber"] = m.group(1)
    else:
        # Try first 9-10 digit number near top
        top = "\n".join(lines[:20])
        for match in _fallback_number_re.finditer(top):
            candidate = match.group(1)
            # skip years like 2021
            if not candidate.startswith("20") or int(candidate) > 2100:
                result["studentNumber"] = candidate
                break

    # Partner university — look for the line before İKÜ
    m = _university_line_re.search(text)
    if m:
        result["partnerUniversity"] = m.group(1).strip().rstrip(",").strip()
    else:
        # Try to find "İstanbul Kültür Üniversitesi" and use the line's left side
        iku_re = re.compile(r'(İstanbul\s*Kültür\s*Üniversitesi|İKÜ)')
        m_iku = iku_re.search(text)
        if m_iku:
            for line in lines:
                if "stanbul" in line and ("K lt r" in line or "Kültür" in line):
                    left = line.split("  ")[0].strip()
                    if left and len(left) > 3:
                        result["partnerUniversity"] = left
                    break

    # Course rows — try full match first (with homeCode), then no_home (direct/bare)
    for line in lines:
        line = line.strip()
        if not line or len(line) < 15:
            continue
        m = _row_full_re.match(line)
        if m:
            result["rows"].append({
                "partnerCode": m.group(1).strip(),
                "partnerName": m.group(2).strip(),
                "localGrade": m.group(3).strip(),
                "ectsGrade": m.group(4).strip(),
                "partnerEcts": m.group(5).strip(),
                "homeCode": m.group(6).strip(),
                "homeName": m.group(7).strip(),
                "ikuGrade": m.group(9).strip(),
                "homeEcts": m.group(10).strip(),
            })
            continue
        m = _row_no_home_re.match(line)
        if m:
            pcode = m.group(1).strip() if m.group(1) else ""
            result["rows"].append({
                "partnerCode": pcode,
                "partnerName": m.group(2).strip(),
                "localGrade": m.group(3).strip(),
                "ectsGrade": m.group(4).strip(),
                "partnerEcts": m.group(5).strip(),
                "homeCode": "",
                "homeName": "",
                "ikuGrade": m.group(7).strip(),
                "homeEcts": m.group(8).strip(),
            })

    # Confidence: need at least student_name OR student_number, and at least 1 row
    has_identity = bool(result["studentName"]) or bool(result["studentNumber"])
    has_rows = len(result["rows"]) >= 1

    if not has_identity or not has_rows:
        return None

    return result


# ── LLM-based parsing ─────────────────────────────────────────────────────────

def _parse_with_llm(text: str) -> dict | None:
    """Fallback: parse via LLM when rule-based fails."""
    prompt = load_prompt("transfer_document_extraction.txt")
    limited = text[:LLM_TEXT_LIMIT]
    prompt = prompt.replace("{text}", limited)

    response = parsing_llm_client.invoke_with_retry(
        prompt,
        min_interval=25,
        max_retries=10,
        start_strict=True,
        expect_json=True,
        context="TransferDoc",
    )

    if not response:
        return None

    extracted = parsing_llm_client.extract_json(response, required_key="rows")
    if not extracted:
        logger.warning("LLM parse produced no valid JSON with 'rows' key")
        return None

    return {
        "studentName": extracted.get("studentName", ""),
        "studentNumber": extracted.get("studentNumber", ""),
        "partnerUniversity": extracted.get("partnerUniversity", ""),
        "homeUniversity": extracted.get("homeUniversity", "İstanbul Kültür Üniversitesi"),
        "rows": extracted.get("rows", []),
    }


# ── Grade conversion ──────────────────────────────────────────────────────────

def _convert_grades(
    db: Session,
    rows: list[dict],
    partner_university_id: int,
) -> tuple[list[dict], list[str], int | None, int | None, int | None, dict | None, list[dict] | None]:
    """Convert grades using the latest grading scheme + ECTS-IKU snapshots.

    Returns (rows, warnings, grading_scheme_id, scheme_version_id, ects_iku_version_id, rules_snapshot, mappings).
    """
    warnings = []
    rules_snapshot = None
    mappings = None
    gs_vid = None
    eiku_vid = None

    scheme = db.query(GradingScheme).filter(
        GradingScheme.university_id == partner_university_id,
        GradingScheme.is_active == True,
    ).first()

    if not scheme:
        warnings.append("No active grading scheme found for the selected university; grades shown as-parsed")
        return rows, warnings, None, None, None, None, None

    gs_vid = get_latest_scheme_snapshot_id(db, scheme.id)
    eiku_vid = get_latest_ects_iku_snapshot_id(db)

    if gs_vid:
        snap = db.query(GradingSchemeSnapshot).filter(GradingSchemeSnapshot.id == gs_vid).first()
        if snap:
            rules_snapshot = snap.rules_snapshot

    if eiku_vid:
        eiku = db.query(EctsIkuSnapshot).filter(EctsIkuSnapshot.id == eiku_vid).first()
        if eiku:
            mappings = eiku.mappings_snapshot

    if not rules_snapshot or not mappings:
        warnings.append("Grading scheme snapshot data incomplete; grades shown as-parsed")
        return rows, warnings, scheme.id, gs_vid, eiku_vid, rules_snapshot, mappings

    converted = 0
    failed = 0
    for row in rows:
        local_grade = row.get("localGrade", "").strip()
        if not local_grade:
            continue
        try:
            result = convert_single_grade_from_snapshot(
                local_grade, rules_snapshot, mappings, has_ects=False,
            )
            if not row.get("ectsGrade"):
                row["ectsGrade"] = result.get("ects_grade", "")
            if not row.get("ikuGrade"):
                row["ikuGrade"] = result.get("iku_grade", "")
            converted += 1
        except ValueError:
            # Try as ECTS grade
            ects_grade = row.get("ectsGrade", "").strip()
            if ects_grade:
                try:
                    result = convert_single_grade_from_snapshot(
                        ects_grade, rules_snapshot, mappings, has_ects=True,
                    )
                    if not row.get("ikuGrade"):
                        row["ikuGrade"] = result.get("iku_grade", "")
                    converted += 1
                except ValueError:
                    failed += 1
            else:
                failed += 1

    if failed > 0:
        warnings.append(f"{failed} grade(s) could not be converted using the grading scheme")

    return rows, warnings, scheme.id, gs_vid, eiku_vid, rules_snapshot, mappings


# ── Main entry point ──────────────────────────────────────────────────────────

def parse_transfer_document(
    db: Session,
    file_path: str,
    partner_university_id: int,
) -> dict:
    """Parse a transfer document end-to-end.

    Returns dict matching the frontend's expected UploadedDocument shape.
    """
    warnings = []

    # 1. Extract text (OCR fallback if text too short)
    text = extract_text_from_file(file_path)
    if len(text.strip()) < MIN_TEXT_LENGTH:
        warnings.append("Document contains very little selectable text; attempting OCR")
        text = _ocr_file(file_path)
        if len(text.strip()) < MIN_TEXT_LENGTH:
            warnings.append("OCR extraction also produced too little text")

    # 2. Parse
    parsing_method = "rule_based"
    result = _parse_rule_based(text)

    if result is None:
        parsing_method = "llm_fallback"
        warnings.append("Rule-based parsing insufficient; using AI extraction")
        result = _parse_with_llm(text)
        if result is None:
            # Return empty but structured
            result = {
                "studentName": "",
                "studentNumber": "",
                "partnerUniversity": "",
                "homeUniversity": "İstanbul Kültür Üniversitesi",
                "rows": [],
            }
            warnings.append("AI extraction also failed; manual review required")

    # 3. Resolve partner university name from DB
    uni = db.query(University).filter(University.id == partner_university_id).first()
    if uni:
        result["partnerUniversity"] = uni.name
    else:
        warnings.append("Selected partner university not found in database")

    # 4. Grade conversion
    rows = result.get("rows", [])
    converted_rows, conv_warnings, gs_id, gs_vid, eiku_vid, rules_snap, mappings = _convert_grades(
        db, rows, partner_university_id,
    )
    result["rows"] = converted_rows
    warnings.extend(conv_warnings)

    result["warnings"] = warnings
    result["parsing_method"] = parsing_method
    result["grading_scheme_id"] = gs_id
    result["_gs_vid"] = gs_vid
    result["_eiku_vid"] = eiku_vid
    result["rules_snapshot"] = rules_snap
    result["mappings"] = mappings

    return result


# ── Verification ──────────────────────────────────────────────────────────────

def _describe_rule(row: dict, rules_snapshot: list[dict]) -> str:
    """Describe which grading rule matched a row's local grade."""
    local = row.get("localGrade", "").strip()
    if not local or not rules_snapshot:
        return ""

    try:
        val = float(local)
        is_numeric = True
    except ValueError:
        is_numeric = False

    for rule in rules_snapshot:
        if is_numeric and not rule.get("local_grade_exact"):
            lo_str = rule.get("local_grade_min")
            hi_str = rule.get("local_grade_max")
            if lo_str is not None and hi_str is not None:
                try:
                    if float(lo_str) <= val <= float(hi_str):
                        return f"{lo_str}–{hi_str} → {rule.get('ects_grade')} ({rule.get('description', '')})".strip()
                except ValueError:
                    continue
        elif not is_numeric and rule.get("local_grade_exact"):
            if rule["local_grade_exact"].strip().lower() == local.lower():
                return f"'{rule['local_grade_exact']}' → {rule.get('ects_grade')} ({rule.get('description', '')})".strip()

    return ""


def verify_document_grades(
    db: Session,
    doc_id: int,
    rows: list[dict],
    rules_snapshot: list[dict] | None,
    mappings: list[dict] | None,
    partner_university_name: str,
    version_number: int = 1,
) -> None:
    """Verify parsed grades against the grading scheme. Writes results to DB."""
    if not rules_snapshot or not mappings:
        return

    valid = invalid = manual = partial = 0

    for i, row in enumerate(rows):
        local_grade = row.get("localGrade", "").strip()
        ects_grade = row.get("ectsGrade", "").strip()
        iku_grade = row.get("ikuGrade", "").strip()

        expected_ects = ""
        expected_iku = ""
        grade_rule = ""
        validation = "no_rule_found"

        if local_grade:
            # Try local→ECTS→IKU first
            local_ects_ok = False
            local_iku_from_local_ects = ""
            try:
                converted = convert_single_grade_from_snapshot(
                    local_grade, rules_snapshot, mappings, has_ects=False,
                )
                expected_ects = converted.get("ects_grade", "")
                local_iku_from_local_ects = converted.get("iku_grade", "")
                grade_rule = _describe_rule(row, rules_snapshot)

                if ects_grade and ects_grade.upper() == expected_ects.upper():
                    # Full match: local→ECTS correct
                    local_ects_ok = True
                    expected_iku = local_iku_from_local_ects
                    
                    if iku_grade and iku_grade.upper() == expected_iku.upper():
                        validation = "valid"
                        valid += 1
                    else:
                        validation = "invalid"
                        invalid += 1
                elif expected_ects:
                    # Local→ECTS produced a result but doesn't match document's ECTS
                    # Check if document's ECTS→IKU matches — could be partial
                    local_ects_ok = False
            except ValueError:
                # Local grade doesn't match any rule
                pass

            if not local_ects_ok and validation == "no_rule_found":
                # Local conversion failed or mismatched — try direct ECTS→IKU
                if ects_grade:
                    try:
                        converted = convert_single_grade_from_snapshot(
                            ects_grade, rules_snapshot, mappings, has_ects=True,
                        )
                        direct_ects_iku = converted.get("iku_grade", "")
                        if local_grade and expected_ects:
                            # Local gave a result but wrong ECTS — ECTS→IKU may still be correct
                            if iku_grade and direct_ects_iku.upper() == iku_grade.upper():
                                validation = "partial"
                                expected_ects = ects_grade
                                expected_iku = direct_ects_iku
                                grade_rule = (grade_rule or "") + f" | Direct ECTS→IKU: {ects_grade}→{direct_ects_iku}"
                                partial += 1
                            else:
                                validation = "invalid"
                                expected_ects = ects_grade
                                expected_iku = direct_ects_iku
                                invalid += 1
                        else:
                            # Local grade invalid/unconvertible, ECTS→IKU is our only reference
                            if iku_grade and direct_ects_iku.upper() == iku_grade.upper():
                                validation = "partial"
                                expected_ects = ects_grade
                                expected_iku = direct_ects_iku
                                grade_rule = f"Direct ECTS: {ects_grade}→{direct_ects_iku}"
                                partial += 1
                            else:
                                validation = "invalid"
                                expected_ects = ects_grade
                                expected_iku = direct_ects_iku
                                grade_rule = f"Direct ECTS: {ects_grade}→{direct_ects_iku}"
                                invalid += 1
                    except ValueError:
                        validation = "no_rule_found"
                        manual += 1
                else:
                    if expected_ects:
                        validation = "invalid"
                        invalid += 1
                    else:
                        validation = "no_rule_found"
                        manual += 1
        else:
            # No local grade — rely on ECTS→IKU only
            if ects_grade:
                try:
                    converted = convert_single_grade_from_snapshot(
                        ects_grade, rules_snapshot, mappings, has_ects=True,
                    )
                    expected_ects = ects_grade
                    expected_iku = converted.get("iku_grade", "")
                    grade_rule = f"Direct ECTS: {ects_grade}→{expected_iku}"

                    if iku_grade and expected_iku.upper() == iku_grade.upper():
                        validation = "partial"
                        partial += 1
                    else:
                        validation = "invalid"
                        invalid += 1
                except ValueError:
                    validation = "no_rule_found"
                    manual += 1
            else:
                validation = "no_rule_found"
                manual += 1

        vr = TransferVerificationResult(
            transfer_document_id=doc_id,
            row_index=i,
            version_number=version_number,
            is_active=True,
            partner_course_name=row.get("partnerName", "")[:500],
            partner_course_code=row.get("partnerCode", "")[:50],
            partner_grade=local_grade[:20],
            partner_ects=str(row.get("partnerEcts", ""))[:10],
            expected_ects_grade=expected_ects[:20],
            expected_iku_grade=expected_iku[:20],
            provided_ects_grade=ects_grade[:20],
            provided_iku_grade=iku_grade[:20],
            validation_result=validation,
            grade_rule_used=grade_rule,
        )
        db.add(vr)

    doc = db.query(TransferDocument).filter(TransferDocument.id == doc_id).first()
    if doc:
        doc.verification_status = "has_issues" if invalid > 0 or manual > 0 or partial > 0 else "verified"
        doc.total_rows = len(rows)
        doc.valid_rows = valid
        doc.invalid_rows = invalid
        doc.partial_rows = partial
        doc.manual_check_rows = manual

    db.flush()


# ── Error explanations ────────────────────────────────────────────────────────

def generate_error_explanations(
    db: Session,
    doc_id: int,
    partner_university_name: str,
) -> None:
    """Generate LLM explanations for invalid/manual_check rows."""
    vrs = db.query(TransferVerificationResult).filter(
        TransferVerificationResult.transfer_document_id == doc_id,
        TransferVerificationResult.validation_result.in_(["invalid", "partial", "manual_check_required", "no_rule_found"]),
    ).all()

    if not vrs:
        return

    prompt_template = load_prompt("transfer_error_explanation.txt")

    for vr in vrs:
        # Build grading scheme summary for context
        scheme_summary = ""
        if vr.transfer_document:
            doc = vr.transfer_document
            if doc.grading_scheme_version_id:
                gs = db.query(GradingSchemeSnapshot).filter(GradingSchemeSnapshot.id == doc.grading_scheme_version_id).first()
                if gs and gs.rules_snapshot:
                    rules = gs.rules_snapshot
                    rule_lines = []
                    for r in rules[:20]:  # limit to 20 rules
                        parts = []
                        if r.get("local_grade_exact"):
                            parts.append(f"'{r['local_grade_exact']}'")
                        elif r.get("local_grade_min") is not None and r.get("local_grade_max") is not None:
                            parts.append(f"{r['local_grade_min']}-{r['local_grade_max']}")
                        if r.get("ects_grade"):
                            parts.append(f"→ {r['ects_grade']}")
                        if r.get("description"):
                            parts.append(f"({r['description']})")
                        rule_lines.append(" ".join(parts))
                    scheme_summary = "\n".join(rule_lines)

        prompt = prompt_template.replace("{university}", partner_university_name)
        prompt = prompt.replace("{course_name}", vr.partner_course_name or "(bilinmiyor)")
        prompt = prompt.replace("{course_code}", vr.partner_course_code or "")
        prompt = prompt.replace("{provided_grade}", f"Yerel Not: {vr.partner_grade}, ECTS: {vr.provided_ects_grade}, İKÜ: {vr.provided_iku_grade}")
        prompt = prompt.replace("{expected_grade}", f"ECTS: {vr.expected_ects_grade}, İKÜ: {vr.expected_iku_grade}")
        prompt = prompt.replace("{grade_rule}", vr.grade_rule_used or "Kural bulunamadı")
        prompt = prompt.replace("{validation_type}", vr.validation_result or "invalid")
        prompt = prompt.replace("{grading_scheme_rules}", scheme_summary or "Şema detayı bulunamadı")

        response = parsing_llm_client.invoke_with_retry(
            prompt,
            min_interval=5,
            max_retries=10,
            start_strict=True,
            expect_json=False,
            context="TransferExplain",
        )

        if response:
            explanation = response.strip()
            # Output guard: filter empty or hallucinating responses
            if len(explanation) > 10 and len(explanation) < 800:
                vr.explanation = explanation
            else:
                vr.explanation = "Açıklama üretilemedi — lütfen manuel kontrol edin."

    db.flush()
